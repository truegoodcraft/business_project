from __future__ import annotations

import io
import os
import time
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List

try:  # pragma: no cover - optional dependency
    from pypdf import PdfReader, PdfWriter
except Exception:  # pragma: no cover - gracefully degrade
    PdfReader = PdfWriter = None  # type: ignore

try:  # pragma: no cover - optional dependency
    from docx import Document
except Exception:  # pragma: no cover - gracefully degrade
    Document = None  # type: ignore


MAX_FILE_SIZE = 50 * 1024 * 1024
MAX_TOTAL_SIZE = 500 * 1024 * 1024
TIME_LIMIT_SECONDS = 10


class BundleBuildError(Exception):
    pass


def _check_limits(paths: Iterable[Path]) -> None:
    total = 0
    for path in paths:
        size = path.stat().st_size
        if size > MAX_FILE_SIZE:
            raise BundleBuildError(f"file_too_large:{path.name}")
        total += size
        if total > MAX_TOTAL_SIZE:
            raise BundleBuildError("bundle_too_large")


def _build_pdf(input_paths: List[Path], out_path: Path) -> Dict[str, object]:
    if PdfReader is None or PdfWriter is None:
        raise BundleBuildError("pdf_support_unavailable")
    writer = PdfWriter()
    pages = 0
    for path in input_paths:
        reader = PdfReader(str(path))
        for page in reader.pages:
            writer.add_page(page)
            pages += 1
    with open(out_path, "wb") as handle:
        writer.write(handle)
    return {"ok": True, "bytes_written": out_path.stat().st_size, "pages": pages, "warnings": []}


def _build_docx(input_paths: List[Path], out_path: Path) -> Dict[str, object]:
    if Document is None:
        raise BundleBuildError("docx_support_unavailable")
    if not input_paths:
        doc = Document()
        doc.add_paragraph("Empty bundle")
        doc.save(out_path)
        return {"ok": True, "bytes_written": out_path.stat().st_size, "warnings": []}
    result = Document(str(input_paths[0]))
    for path in input_paths[1:]:
        segment = Document(str(path))
        for element in segment.element.body:
            result.element.body.append(element)
    result.save(out_path)
    return {"ok": True, "bytes_written": out_path.stat().st_size, "warnings": []}


def _build_text(input_paths: List[Path], out_path: Path) -> Dict[str, object]:
    buffer = io.StringIO()
    warnings: List[str] = []
    for idx, path in enumerate(input_paths):
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            warnings.append(f"non_utf8:{path.name}")
            text = path.read_text(encoding="utf-8", errors="replace")
        if idx:
            buffer.write("\n\n")
        buffer.write(text)
    data = buffer.getvalue().encode("utf-8")
    with open(out_path, "wb") as handle:
        handle.write(data)
    return {"ok": True, "bytes_written": len(data), "warnings": warnings}


def _build_zip(input_paths: List[Path], out_path: Path) -> Dict[str, object]:
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in input_paths:
            archive.write(path, arcname=path.name)
    return {"ok": True, "bytes_written": out_path.stat().st_size, "warnings": []}


def build_bundle(mode: str, input_paths: Iterable[str], out_tmp_path: Path) -> Dict[str, object]:
    paths = [Path(p) for p in input_paths]
    start = time.time()
    _check_limits(paths)
    for path in paths:
        if not path.exists():
            raise BundleBuildError(f"missing_input:{path}")
    if time.time() - start > TIME_LIMIT_SECONDS:
        raise BundleBuildError("time_budget_exceeded")
    tmp_parent = out_tmp_path.parent
    tmp_parent.mkdir(parents=True, exist_ok=True)

    mode = mode or "zip_bundle"
    if mode == "pdf_merge":
        result = _build_pdf(paths, out_tmp_path)
    elif mode == "docx_merge":
        result = _build_docx(paths, out_tmp_path)
    elif mode == "text_concat":
        result = _build_text(paths, out_tmp_path)
    else:
        result = _build_zip(paths, out_tmp_path)
    if time.time() - start > TIME_LIMIT_SECONDS:
        raise BundleBuildError("time_budget_exceeded")
    result.setdefault("warnings", [])
    result.setdefault("bytes_written", out_tmp_path.stat().st_size)
    return result


__all__ = ["build_bundle", "BundleBuildError"]
